"""
Генерация Word-отчёта по проекту классификации ЭКГ (MIT-BIH).
Запуск:  pip install python-docx  &&  python generate_report.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def build():
    doc = Document()

    # Базовый стиль
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ===== Титул =====
    title = doc.add_heading("Нейросетевая классификация патологий ритма сердца", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Отчёт по проекту обработки и классификации сигналов ЭКГ (база MIT-BIH)")
    r.italic = True
    r.font.size = Pt(13)
    doc.add_paragraph()

    # ===== 1. Тема, описание задачи =====
    add_heading(doc, "1. Тема, описание задачи", 1)
    add_para(doc,
        "Тема работы — автоматическая классификация отдельных сердечных сокращений "
        "(сегментов ЭКГ) на 5 классов по стандарту AAMI с помощью свёрточной нейронной сети.")
    add_para(doc, "Классы AAMI:", bold=True)
    make_table(doc,
        ["Код", "Название", "Описание"],
        [
            ["N", "Normal", "Нормальные сокращения"],
            ["S", "Supraventricular ectopic", "Наджелудочковые экстрасистолы"],
            ["V", "Ventricular ectopic", "Желудочковые экстрасистолы"],
            ["F", "Fusion", "Сливные комплексы"],
            ["Q", "Unknown", "Неизвестные / артефакты"],
        ])
    doc.add_paragraph()
    add_para(doc, "Цель и подзадачи:", bold=True)
    add_numbered(doc, [
        "Предобработка сигналов ЭКГ: фильтрация, нормализация, сегментация вокруг R-пиков.",
        "Устранение дисбаланса классов (класс N составляет около 83% данных).",
        "Построение и обучение 1D-CNN с residual-блоками.",
        "Оценка качества модели и анализ ошибок по классам.",
    ])
    add_para(doc,
        "Практическая значимость: автоматический скрининг аритмий позволяет ускорить "
        "анализ длительных записей ЭКГ и снизить нагрузку на врача-кардиолога.")

    # ===== 2. База =====
    add_heading(doc, "2. База", 1)
    add_para(doc,
        "В работе используется база MIT-BIH Arrhythmia Database — стандартный набор данных "
        "для исследований по классификации аритмий.")
    add_para(doc, "Характеристики базы:", bold=True)
    add_bullets(doc, [
        "Частота дискретизации: 360 Гц.",
        "Количество записей: 48 (по два отведения, используется первое — обычно MLII).",
        "Перевод из ADC в милливольты: gain = 200 ADU/мВ, base = 1024.",
        "Формат: CSV (сигнал) + TXT (аннотации R-пиков и типов сокращений).",
        "Всего сегментов после предобработки: 109 449.",
    ])
    doc.add_paragraph()
    add_para(doc, "Исходное распределение классов (все данные):", bold=True)
    make_table(doc,
        ["Класс", "Количество", "Доля, %"],
        [
            ["N", "90 592", "82.77"],
            ["S", "2 781", "2.54"],
            ["V", "7 235", "6.61"],
            ["F", "802", "0.73"],
            ["Q", "8 039", "7.34"],
            ["Всего", "109 449", "100.00"],
        ])
    add_para(doc,
        "Коэффициент дисбаланса (max/min) составляет около 113:1, что делает метрику accuracy "
        "неинформативной и требует применения методов балансировки.", italic=True)

    # ===== 3. Параметризация данных =====
    add_heading(doc, "3. Параметризация данных", 1)
    add_para(doc, "3.1. Предобработка сигнала", bold=True)
    add_numbered(doc, [
        "Полосовая фильтрация Баттерворта 0.5–40 Гц (порядок 4) — удаление дрейфа изолинии "
        "и высокочастотного шума.",
        "Режекторный фильтр (notch) 50 Гц — подавление сетевой наводки.",
        "Z-нормализация сигнала: приведение к среднему 0 и стандартному отклонению 1.",
        "Сегментация: окно [R-100, R+150] = 250 отсчётов (~694 мс) вокруг каждого R-пика.",
    ])
    doc.add_paragraph()
    add_para(doc, "3.2. Разделение на train/test", bold=True)
    add_para(doc,
        "Применяется межпациентское разделение (DS1/DS2): записи одного пациента не попадают "
        "одновременно в обучающую и тестовую выборки. Это исключает «утечку» данных и даёт "
        "честную оценку обобщающей способности.")
    make_table(doc,
        ["Выборка", "Размер", "Балансировка"],
        [
            ["Train (DS1)", "51 001", "Применяется"],
            ["Test (DS2)", "49 691", "НЕ применяется"],
        ])
    doc.add_paragraph()
    add_para(doc, "3.3. Методы устранения дисбаланса классов", bold=True)
    add_para(doc, "Подготовлено три варианта обучающей выборки:")
    add_numbered(doc, [
        "Weighted — исходные несбалансированные данные + взвешенная функция потерь "
        "CrossEntropyLoss с весами классов.",
        "SMOTE — синтетическая интерполяция миноритарных классов; объём вырос до 137 539 примеров.",
        "Аугментации (рекомендуется для ЭКГ) — физически осмысленные искажения сигнала: "
        "гауссов шум, масштабирование амплитуды, временной сдвиг, дрейф изолинии, time warping; "
        "объём 106 975 примеров.",
    ])
    doc.add_paragraph()
    add_para(doc, "Веса классов для взвешенного loss:", bold=True)
    make_table(doc,
        ["Класс", "Вес"],
        [
            ["N", "0.001"],
            ["S", "0.041"],
            ["V", "0.010"],
            ["F", "0.094"],
            ["Q", "4.854"],
        ])

    # ===== 4. Архитектура нейросети =====
    add_heading(doc, "4. Архитектура нейросети", 1)
    add_para(doc,
        "Используется одномерная свёрточная сеть с residual-блоками (1D ResNet), "
        "адаптированная под задачу классификации ЭКГ-сегментов.")
    add_para(doc, "Структура модели (ECGResNet):", bold=True)
    add_numbered(doc, [
        "Stem — начальный свёрточный слой Conv1D (kernel=7) + BatchNorm + ReLU + MaxPool "
        "(250 → 125 отсчётов).",
        "Residual-блок 1: 32 → 32 канала (stride=1).",
        "Residual-блок 2: 32 → 64 канала (stride=2, 125 → 63).",
        "Residual-блок 3: 64 → 128 каналов (stride=2, 63 → 32).",
        "Global Average Pooling — агрегация по временной оси.",
        "Классификатор: Dropout(0.3) → Linear(128→64) → ReLU → Dropout(0.3) → Linear(64→5).",
    ])
    doc.add_paragraph()
    add_para(doc, "Параметры residual-блока:", bold=True)
    add_bullets(doc, [
        "Два свёрточных слоя Conv1D (kernel=5) с BatchNorm и ReLU.",
        "Skip-соединение y = F(x) + x для стабилизации градиентов.",
        "Dropout 0.2–0.3 для регуляризации.",
    ])
    doc.add_paragraph()
    make_table(doc,
        ["Параметр", "Значение"],
        [
            ["Тип сети", "1D ResNet (ECGResNet)"],
            ["Базовое число каналов", "32"],
            ["Dropout", "0.3"],
            ["Число обучаемых параметров", "184 229"],
            ["Вход", "(batch, 1, 250)"],
            ["Выход", "(batch, 5)"],
        ])
    doc.add_paragraph()
    add_para(doc, "Параметры обучения:", bold=True)
    make_table(doc,
        ["Параметр", "Значение"],
        [
            ["Функция потерь", "CrossEntropyLoss"],
            ["Оптимизатор", "Adam (lr=1e-3, weight_decay=1e-4)"],
            ["Scheduler", "ReduceLROnPlateau (factor=0.5, patience=3)"],
            ["Batch size", "256"],
            ["Эпох (макс.)", "30"],
            ["Early Stopping", "patience=7 по macro-F1"],
            ["Стратегия балансировки", "aug (аугментации)"],
        ])

    # ===== 5. Графическое подтверждение =====
    add_heading(doc, "5. Графическое подтверждение", 1)
    add_para(doc,
        "Ниже приведена динамика обучения по эпохам. На валидации модель достигла "
        "macro-F1 = 0.9944. Полные графики (loss, accuracy, macro-F1, confusion matrix) "
        "доступны в ноутбуке BuildingModel.ipynb.")
    add_para(doc, "Динамика метрик по ключевым эпохам:", bold=True)
    make_table(doc,
        ["Эпоха", "Train Loss", "Train F1", "Val Loss", "Val F1"],
        [
            ["1", "0.3413", "0.8731", "0.1906", "0.9426"],
            ["10", "0.0517", "0.9817", "0.0337", "0.9888"],
            ["18", "0.0371", "0.9869", "0.0221", "0.9925"],
            ["24", "0.0309", "0.9889", "0.0198", "0.9925"],
            ["29 (лучшая)", "0.0199", "0.9932", "0.0156", "0.9944"],
            ["30", "0.0185", "0.9938", "0.0166", "0.9942"],
        ])
    doc.add_paragraph()
    # Вставка изображений, если они сохранены пользователем
    for img_name, caption in [
        ("training_curves.png", "Рис. 1. Кривые обучения: loss, accuracy, macro-F1."),
        ("confusion_matrix.png", "Рис. 2. Матрица ошибок на тестовой выборке."),
    ]:
        try:
            doc.add_picture(img_name, width=Inches(6.0))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.italic = True
        except Exception:
            note = doc.add_paragraph(
                f"[Здесь вставьте изображение «{img_name}», экспортированное из ноутбука. "
                f"Подпись: {caption}]")
            for run in note.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()
    add_para(doc, "Метрики на тестовой выборке (по классам):", bold=True)
    make_table(doc,
        ["Класс", "Precision", "Recall", "F1", "Support"],
        [
            ["N", "0.9539", "0.8644", "0.9069", "44 239"],
            ["S", "0.0468", "0.0947", "0.0627", "1 837"],
            ["V", "0.5486", "0.9391", "0.6926", "3 220"],
            ["F", "0.0134", "0.0129", "0.0132", "388"],
            ["Q", "0.0000", "0.0000", "0.0000", "7"],
            ["Accuracy", "", "", "0.8340", "49 691"],
            ["Macro avg", "0.3126", "0.3822", "0.3351", "49 691"],
        ])

    # ===== 6. Выводы =====
    add_heading(doc, "6. Выводы", 1)
    add_numbered(doc, [
        "Реализован полный пайплайн обработки ЭКГ: фильтрация, нормализация, сегментация "
        "и классификация на 5 классов AAMI.",
        "На валидации (часть DS1) достигнут высокий macro-F1 = 0.9944, что подтверждает "
        "способность модели обучаться на сбалансированных данных.",
        "На независимой тестовой выборке (DS2, межпациентское разбиение) Accuracy = 0.8340, "
        "Macro-F1 = 0.3351. Заметное снижение Macro-F1 указывает на проблему обобщения "
        "на новых пациентов, особенно для миноритарных классов S, F, Q.",
        "Класс N распознаётся надёжно (F1 = 0.91), класс V — удовлетворительно (F1 = 0.69), "
        "однако классы S, F и Q практически не распознаются на тесте.",
        "Большой разрыв между валидацией и тестом — следствие межпациентской вариабельности "
        "морфологии ЭКГ и сильного дисбаланса: аугментации/SMOTE улучшают баланс внутри "
        "DS1, но не покрывают разнообразие пациентов DS2.",
    ])

    # ===== 7. План дальнейшей работы =====
    add_heading(doc, "7. План дальнейшей работы", 1)
    add_numbered(doc, [
        "Сравнить три стратегии балансировки (weighted, SMOTE, aug) на тесте DS2 и выбрать "
        "оптимальную по macro-F1.",
        "Добавить контекст соседних сокращений (RR-интервалы, признаки предыдущего/следующего "
        "комплекса) — это особенно важно для класса S.",
        "Применить focal loss вместо CrossEntropyLoss для усиления вклада трудных и "
        "миноритарных классов.",
        "Расширить и усилить аугментации, ориентированные на межпациентскую вариабельность; "
        "рассмотреть domain adaptation.",
        "Исследовать более ёмкие архитектуры: глубокие 1D ResNet, CNN+BiLSTM, "
        "механизмы внимания (attention), трансформеры для временных рядов.",
        "Провести кросс-валидацию по пациентам (k-fold по pid) для устойчивой оценки.",
        "Откалибровать вероятности и оценить модель по клинически значимым метрикам "
        "(чувствительность/специфичность для аритмий).",
        "Подготовить инференс-пайплайн для записей в реальном времени и протестировать "
        "на сторонних базах ЭКГ.",
    ])

    out = "Отчёт_ЭКГ_классификация.docx"
    doc.save(out)
    print(f"Документ сохранён: {out}")


if __name__ == "__main__":
    build()